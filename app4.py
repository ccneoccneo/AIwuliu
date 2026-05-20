import gradio as gr
import tempfile
import os
from agent import chat

# ── 导出功能需要的库 ──
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# ── 1. 主题定制 ──
theme = gr.themes.Soft(
    primary_hue="blue",
    secondary_hue="slate",
    neutral_hue="slate",
).set(
    body_background_fill="*neutral_50",
    block_background_fill="white",
    block_label_background_fill="*primary_100",
    button_primary_background_fill="*primary_600",
    button_primary_background_fill_hover="*primary_700",
    button_primary_text_color="white",
    border_color_primary="*primary_200",
    input_background_fill="white",
)


# ── 2. 自定义 CSS（保持之前的美化） ──
css = """
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&display=swap');

* {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC',
                 'Hiragino Sans GB', 'Microsoft YaHei', sans-serif !important;
}

.gradio-container {
    background: linear-gradient(135deg, #f5f7fa 0%, #e9edf5 100%) !important;
}

.main-card {
    max-width: 880px;
    margin: 24px auto !important;
    border-radius: 24px !important;
    box-shadow: 0 12px 40px rgba(0,0,0,0.06), 0 2px 8px rgba(0,0,0,0.04) !important;
    background: white !important;
    padding: 24px 28px !important;
}

#chatbot {
    border-radius: 16px !important;
    border: 1px solid #e8ecf1 !important;
    background: #fafbfc !important;
    box-shadow: inset 0 1px 3px rgba(0,0,0,0.02) !important;
}

.example {
    border-radius: 20px !important;
    padding: 6px 18px !important;
    font-size: 0.875rem !important;
    font-weight: 500 !important;
    background: #f0f4ff !important;
    border: 1px solid #d0dcff !important;
    color: #1e3a8a !important;
    transition: all 0.15s ease;
}
.example:hover {
    background: #e0e9ff !important;
    border-color: #b0c4ff !important;
    transform: translateY(-1px);
}

#input-box textarea {
    border-radius: 20px !important;
    padding: 12px 18px !important;
    border: 1px solid #dce1e8 !important;
    background: white !important;
    transition: border-color 0.15s ease;
}
#input-box textarea:focus {
    border-color: #4f8cff !important;
    box-shadow: 0 0 0 3px rgba(79,140,255,0.1) !important;
}

#send-btn {
    border-radius: 20px !important;
    font-weight: 600 !important;
    padding: 10px 24px !important;
    background: linear-gradient(135deg, #4f8cff 0%, #1e5ee0 100%) !important;
    border: none !important;
    box-shadow: 0 2px 6px rgba(79,140,255,0.3);
    transition: all 0.2s ease;
}
#send-btn:hover {
    transform: translateY(-1px);
    box-shadow: 0 4px 12px rgba(79,140,255,0.4);
}

.app-title {
    font-size: 2rem;
    font-weight: 700;
    background: linear-gradient(135deg, #1e3a8a, #3b82f6);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    text-align: center;
    margin-bottom: 4px;
}
.app-subtitle {
    font-size: 0.95rem;
    color: #6b7c93;
    text-align: center;
    margin-bottom: 20px;
    font-weight: 400;
}

/* 导出按钮组样式 */
.export-row {
    display: flex;
    gap: 10px;
    margin-top: 10px;
}
"""


# ── 3. 示例问题 ──
examples = [
    "查询订单 ORD202604130001",
    "最近有延误的订单吗？帮我分析原因",
    "我明天有个上海到北京的急件，怎么发最划算？",
    "如何降低20%的运输成本？",
]


# ── 4. 对话处理函数 ──
def respond(message, history):
    history = history or []
    try:
        response = chat(message, user_id="gradio_user")
    except Exception as e:
        response = f"出错了：{e}"
    history.append((message, response))
    return "", history


# ── 5. 导出功能 ──
def find_chinese_font():
    """在常见路径中搜索中文字体文件"""
    possible_paths = [
        # Windows
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        # macOS
        "/System/Library/Fonts/PingFang.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        # Linux
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    for p in possible_paths:
        if os.path.exists(p):
            return p
    return None


def export_to_word(history):
    """将对话历史导出为 Word 文档，返回文件路径"""
    if not history:
        return None

    doc = Document()
    doc.add_heading('物流AI调度助手 对话记录', 0)

    for i, (user_msg, bot_msg) in enumerate(history, 1):
        doc.add_heading(f'第 {i} 轮对话', level=2)
        doc.add_heading('用户：', level=3)
        doc.add_paragraph(user_msg)
        doc.add_heading('助手：', level=3)
        doc.add_paragraph(bot_msg)

    temp_dir = tempfile.gettempdir()
    filepath = os.path.join(temp_dir, "chat_history.docx")
    doc.save(filepath)
    return filepath


def export_to_pdf(history):
    """将对话历史导出为 PDF 文档，返回文件路径"""
    if not history:
        return None

    font_path = find_chinese_font()
    if not font_path:
        raise RuntimeError("未找到系统中文字体，无法生成 PDF。请使用 Word 导出，或手动安装中文字体。")

    # 注册字体
    try:
        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
    except Exception:
        # 可能是已经注册过
        pass

    temp_dir = tempfile.gettempdir()
    filepath = os.path.join(temp_dir, "chat_history.pdf")
    doc = SimpleDocTemplate(filepath, pagesize=A4)

    # 样式
    styles = getSampleStyleSheet()
    style_normal = ParagraphStyle(
        'ChineseNormal',
        parent=styles['Normal'],
        fontName='ChineseFont',
        fontSize=12,
        leading=20,
        spaceAfter=6,
    )
    style_heading = ParagraphStyle(
        'ChineseHeading',
        parent=styles['Heading2'],
        fontName='ChineseFont',
        spaceAfter=8,
        spaceBefore=12,
    )

    elements = []
    elements.append(Paragraph("物流AI调度助手 对话记录", styles['Title']))
    elements.append(Spacer(1, 12))

    for i, (user_msg, bot_msg) in enumerate(history, 1):
        elements.append(Paragraph(f"第 {i} 轮对话", style_heading))
        elements.append(Paragraph("<b>用户：</b>", style_normal))
        elements.append(Paragraph(user_msg.replace('\n', '<br/>'), style_normal))
        elements.append(Paragraph("<b>助手：</b>", style_normal))
        elements.append(Paragraph(bot_msg.replace('\n', '<br/>'), style_normal))
        elements.append(Spacer(1, 8))

    doc.build(elements)
    return filepath


# ── 6. 构建界面 ──
with gr.Blocks(theme=theme, css=css, title="物流智能问答与决策辅助系统") as demo:
    gr.HTML("""
    <div class="app-title">🚚 物流智能问答与决策辅助系统</div>
    <div class="app-subtitle">查订单 · 问延误 · 优化路线 — 智能物流中枢</div>
    """)

    with gr.Column(elem_classes="main-card"):
        chatbot = gr.Chatbot(
            elem_id="chatbot",
            height=520,
            placeholder="在下方输入您的问题…",
            avatar_images=(
                None,
                "https://img.icons8.com/fluency/48/chatbot.png",
            ),
            bubble_full_width=False,
            layout="bubble",
        )

        with gr.Row():
            msg = gr.Textbox(
                show_label=False,
                placeholder="例如：查订单、问延误、优化路线…",
                container=False,
                scale=8,
                elem_id="input-box",
            )
            submit_btn = gr.Button("发送", variant="primary", elem_id="send-btn")

        gr.Examples(
            examples=examples,
            inputs=msg,
            label=None,
        )

        # 功能按钮行：清空 + 导出
        with gr.Row(elem_classes="export-row"):
            clear_btn = gr.Button("🗑️ 清空对话", variant="secondary", size="sm")
            word_btn = gr.Button("📄 导出为 Word", variant="secondary", size="sm")
            pdf_btn = gr.Button("📑 导出为 PDF", variant="secondary", size="sm")

        # 下载文件组件（点击导出后显示）
        download_file = gr.File(label="下载文件", visible=False)

    # 事件绑定
    msg.submit(respond, [msg, chatbot], [msg, chatbot])
    submit_btn.click(respond, [msg, chatbot], [msg, chatbot])

    clear_btn.click(lambda: [], outputs=[chatbot])
    clear_btn.click(lambda: gr.update(visible=False), outputs=[download_file])

    word_btn.click(
        fn=export_to_word,
        inputs=[chatbot],
        outputs=[download_file],
    ).then(
        fn=lambda: gr.update(visible=True),
        outputs=[download_file],
    )

    pdf_btn.click(
        fn=export_to_pdf,
        inputs=[chatbot],
        outputs=[download_file],
    ).then(
        fn=lambda: gr.update(visible=True),
        outputs=[download_file],
    )


if __name__ == "__main__":
    demo.launch(server_name="127.0.0.1")