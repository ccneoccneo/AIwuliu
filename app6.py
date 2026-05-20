import streamlit as st
import tempfile
import os
from agent import chat  # 你的业务逻辑

# 导出需要
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ---------- 页面配置 ----------
st.set_page_config(page_title="物流智能问答与决策辅助系统", layout="wide")

# ---------- 自定义 CSS ----------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&display=swap');

html, body, [class*="st-"] {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC',
                 'Hiragino Sans GB', 'Microsoft YaHei', sans-serif !important;
}

/* 背景 */
.stApp {
    background: linear-gradient(135deg, #f5f7fa 0%, #e9edf5 100%);
}

/* 主卡片 */
.main-card {
    max-width: 880px;
    margin: 24px auto;
    border-radius: 24px;
    padding: 24px 28px;
    background: white;
    box-shadow: 0 12px 40px rgba(0,0,0,0.06), 0 2px 8px rgba(0,0,0,0.04);
}

/* 标题 */
.app-title {
    font-size: 2rem;
    font-weight: 700;
    background: linear-gradient(135deg, #1e3a8a, #3b82f6);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    text-align: center;
    margin-bottom: 4px;
}
.app-subtitle {
    font-size: 0.95rem;
    color: #6b7c93;
    text-align: center;
    margin-bottom: 20px;
    font-weight: 400;
}

/* 聊天消息气泡 */
.stChatMessage {
    border-radius: 12px;
    border: 1px solid #e8ecf1;
    background: #fafbfc;
    box-shadow: inset 0 1px 3px rgba(0,0,0,0.02);
}

/* 按钮 */
div.stButton > button {
    border-radius: 20px;
    font-weight: 600;
    padding: 10px 24px;
    border: none;
    box-shadow: 0 2px 6px rgba(79,140,255,0.3);
    transition: all 0.2s ease;
}
div.stButton > button:hover {
    transform: translateY(-1px);
    box-shadow: 0 4px 12px rgba(79,140,255,0.4);
}

/* 发送按钮特殊样式（主按钮） */
div.stButton > button[kind="primary"] {
    background: linear-gradient(135deg, #4f8cff 0%, #1e5ee0 100%);
    color: white;
}

/* 示例按钮 */
.example-button button {
    background: #f0f4ff !important;
    border: 1px solid #d0dcff !important;
    color: #1e3a8a !important;
    font-weight: 500;
}
.example-button button:hover {
    background: #e0e9ff !important;
    border-color: #b0c4ff !important;
}
</style>
""", unsafe_allow_html=True)

# ---------- 标题 ----------
st.markdown('<div class="app-title">🚚 物流智能问答与决策辅助系统</div>', unsafe_allow_html=True)
st.markdown('<div class="app-subtitle">查订单 · 问延误 · 优化路线 — 智能物流中枢</div>', unsafe_allow_html=True)

# ---------- 初始化会话状态 ----------
if "messages" not in st.session_state:
    st.session_state.messages = []  # 每条消息 {"role": "user"/"assistant", "content": str}

# ---------- 示例问题处理 ----------
EXAMPLES = [
    "查询订单 ORD202604130001",
    "最近有延误的订单吗？帮我分析原因",
    "我明天有个上海到北京的急件，怎么发最划算？",
    "如何降低20%的运输成本？",
]

def send_example(example):
    """将示例问题作为用户消息发送并调用 agent"""
    st.session_state.messages.append({"role": "user", "content": example})
    with st.spinner("思考中…"):
        try:
            reply = chat(example, user_id="streamlit_user")
        except Exception as e:
            reply = f"出错了：{e}"
    st.session_state.messages.append({"role": "assistant", "content": reply})

# ---------- 聊天界面 ----------
with st.container():
    st.markdown('<div class="main-card">', unsafe_allow_html=True)

    # 显示历史消息
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # 输入框
    prompt = st.chat_input("在这里输入你的问题…")

    if prompt:
        # 添加用户消息
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        # 获取助手回复
        with st.chat_message("assistant"):
            with st.spinner("思考中…"):
                try:
                    reply = chat(prompt, user_id="streamlit_user")
                except Exception as e:
                    reply = f"出错了：{e}"
                st.markdown(reply)
        st.session_state.messages.append({"role": "assistant", "content": reply})

    st.markdown('</div>', unsafe_allow_html=True)

# ---------- 操作区：清空 + 导出 ----------
col1, col2, col3, col4, col5 = st.columns([1, 1, 1, 1, 1])

with col1:
    if st.button("🗑️ 清空对话"):
        st.session_state.messages = []
        st.rerun()

# 导出功能
def get_history_tuples():
    """将 st.session_state.messages 转换为 (user, assistant) 元组列表"""
    tuples = []
    for i in range(0, len(st.session_state.messages)-1, 2):
        if st.session_state.messages[i]["role"] == "user" and st.session_state.messages[i+1]["role"] == "assistant":
            tuples.append((st.session_state.messages[i]["content"], st.session_state.messages[i+1]["content"]))
    return tuples

def find_chinese_font():
    """搜索系统中文字体"""
    paths = [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "/System/Library/Fonts/PingFang.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    for p in paths:
        if os.path.exists(p):
            return p
    return None

def export_word(history):
    doc = Document()
    doc.add_heading('物流AI调度助手 对话记录', 0)
    for i, (user_msg, bot_msg) in enumerate(history, 1):
        doc.add_heading(f'第 {i} 轮对话', level=2)
        doc.add_heading('用户：', level=3)
        doc.add_paragraph(user_msg)
        doc.add_heading('助手：', level=3)
        doc.add_paragraph(bot_msg)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name

def export_pdf(history):
    font_path = find_chinese_font()
    if not font_path:
        raise RuntimeError("未找到中文字体，无法生成 PDF，请使用 Word 或 HTML 导出")

    try:
        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
    except:
        pass

    filepath = os.path.join(tempfile.gettempdir(), "chat_history.pdf")
    doc = SimpleDocTemplate(filepath, pagesize=A4)
    styles = getSampleStyleSheet()
    style_normal = ParagraphStyle('ChineseNormal', parent=styles['Normal'],
                                  fontName='ChineseFont', fontSize=12, leading=20, spaceAfter=6)
    style_heading = ParagraphStyle('ChineseHeading', parent=styles['Heading2'],
                                   fontName='ChineseFont', spaceAfter=8, spaceBefore=12)

    elements = [Paragraph("物流AI调度助手 对话记录", styles['Title']), Spacer(1, 12)]
    for i, (user_msg, bot_msg) in enumerate(history, 1):
        elements.append(Paragraph(f"第 {i} 轮对话", style_heading))
        elements.append(Paragraph("<b>用户：</b>", style_normal))
        elements.append(Paragraph(user_msg.replace('\n', '<br/>'), style_normal))
        elements.append(Paragraph("<b>助手：</b>", style_normal))
        elements.append(Paragraph(bot_msg.replace('\n', '<br/>'), style_normal))
        elements.append(Spacer(1, 8))
    doc.build(elements)
    return filepath

def export_html(history):
    html = '''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>物流AI调度助手 对话记录</title>
    <style>
        body { font-family: 'Segoe UI', 'PingFang SC', 'Microsoft YaHei', sans-serif; background: #f5f7fa; padding: 20px; }
        .container { max-width: 800px; margin: auto; background: white; border-radius: 16px; box-shadow: 0 8px 24px rgba(0,0,0,0.08); padding: 30px; }
        h1 { text-align: center; color: #1e3a8a; }
        .chat-round { margin-bottom: 24px; border-bottom: 1px solid #e8ecf1; padding-bottom: 16px; }
        .round-title { font-weight: 600; color: #4f8cff; margin-bottom: 10px; font-size: 16px; }
        .msg { margin: 8px 0 12px 0; padding: 10px 16px; border-radius: 12px; background: #f0f4ff; line-height: 1.6; white-space: pre-wrap; }
        .msg.user { background: #e0f2fe; border-left: 4px solid #3b82f6; }
        .msg.assistant { background: #f0fdf4; border-left: 4px solid #22c55e; }
        .label { font-weight: 600; font-size: 14px; color: #334155; }
        .footer { text-align: center; margin-top: 30px; color: #94a3b8; }
    </style>
</head>
<body>
    <div class="container">
        <h1>🚚 物流AI调度助手 对话记录</h1>
'''
    for i, (user_msg, bot_msg) in enumerate(history, 1):
        html += f'''
        <div class="chat-round">
            <div class="round-title">第 {i} 轮对话</div>
            <div class="label">👤 用户：</div>
            <div class="msg user">{user_msg}</div>
            <div class="label">🤖 助手：</div>
            <div class="msg assistant">{bot_msg}</div>
        </div>'''
    html += '''
        <div class="footer">由物流AI调度助手自动生成</div>
    </div>
</body>
</html>'''
    filepath = os.path.join(tempfile.gettempdir(), "chat_history.html")
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(html)
    return filepath

history = get_history_tuples()
if not history:
    st.caption("还没有对话，无法导出")

with col2:
    if st.button("📄 导出 Word", disabled=len(history)==0):
        with st.spinner("生成 Word…"):
            path = export_word(history)
            with open(path, "rb") as f:
                st.download_button("下载 Word", f, file_name="chat_history.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

with col3:
    if st.button("📑 导出 PDF", disabled=len(history)==0):
        try:
            with st.spinner("生成 PDF…"):
                path = export_pdf(history)
                with open(path, "rb") as f:
                    st.download_button("下载 PDF", f, file_name="chat_history.pdf", mime="application/pdf")
        except RuntimeError as e:
            st.error(str(e))

with col4:
    if st.button("🌐 导出 HTML", disabled=len(history)==0):
        with st.spinner("生成 HTML…"):
            path = export_html(history)
            with open(path, "rb") as f:
                st.download_button("下载 HTML", f, file_name="chat_history.html", mime="text/html")

# ---------- 示例问题快速入口 ----------
st.markdown("---")
st.markdown("**💡 试试这些示例：**")
cols = st.columns(4)
for idx, example in enumerate(EXAMPLES):
    with cols[idx]:
        if st.button(example, key=f"ex_{idx}", use_container_width=True, help="点击发送这条示例"):
            send_example(example)
            st.rerun()